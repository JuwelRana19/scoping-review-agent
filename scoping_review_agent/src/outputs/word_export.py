from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

from scoping_review_agent.src.utils.io import ensure_dir, read_jsonl


def _add_field(doc: Any, field_name: str, value: Any, evidence_quotes: Dict[str, Any]) -> None:
    doc.add_heading(field_name.replace("_", " ").title(), level=3)
    val = "" if value is None else str(value)
    if not val.strip():
        doc.add_paragraph("(Not extracted / insufficient evidence)")
        return
    doc.add_paragraph(val)

    # Evidence quotes
    quotes = evidence_quotes.get(field_name, []) if evidence_quotes else []
    if isinstance(quotes, list) and quotes:
        doc.add_paragraph("Evidence quotes (excerpts):")
        for q in quotes[:5]:
            if not isinstance(q, dict):
                continue
            quote = q.get("quote") or ""
            locator = q.get("citation_locator") or ""
            if not quote:
                continue
            p = doc.add_paragraph(quote)
            p.runs[0].italic = True
            if locator:
                doc.add_paragraph(f"Locator: {locator}")


def export_word_document(
    *,
    extractions_jsonl_path: str | Path,
    output_dir: str | Path,
    title: str = "Scoping review extraction (human review)",
) -> Path:
    from docx import Document

    output_dir = Path(output_dir)
    ensure_dir(output_dir)
    rows = read_jsonl(extractions_jsonl_path)

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("Editable fields are intended for human verification and correction.")

    for row in rows:
        paper_heading = (row.get("title") or "Untitled").strip()
        paper_id = str(row.get("paper_id") or "").strip()
        doc.add_page_break()
        doc.add_heading(paper_heading, level=2)

        if paper_id:
            doc.add_paragraph("Paper ID: " + paper_id)

        if row.get("authors"):
            doc.add_paragraph("Authors: " + ", ".join([str(a) for a in row.get("authors")][:12]))

        if row.get("pmid"):
            doc.add_paragraph("PMID: " + str(row.get("pmid")))
        if row.get("doi"):
            doc.add_paragraph("DOI: https://doi.org/" + str(row.get("doi")))

        # Human section
        doc.add_heading("Human review", level=3)
        doc.add_paragraph("Decision (include/exclude/uncertain): ")
        doc.add_paragraph("Notes / corrections: ")

        evidence_quotes = row.get("evidence_quotes") or {}
        if not isinstance(evidence_quotes, dict):
            evidence_quotes = {}

        required_fields = [
            "study_design",
            "research_objective",
            "causal_or_epidemiologic_methods",
            "key_findings",
            "strengths",
            "limitations",
            "future_research_agenda",
        ]
        for f in required_fields:
            _add_field(doc, f, row.get(f), evidence_quotes)

    out_path = output_dir / "extraction_human_review.docx"
    doc.save(out_path)
    return out_path

